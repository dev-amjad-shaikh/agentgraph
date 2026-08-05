#!/usr/bin/env python3
"""Markdown -> DOCX converter for the agentgraph blog post.

Adapted from whitepaper/md2docx.py.

Handles:
  - # / ## / ### headings -> Word heading styles
  - **bold**, *italic*, `code` inline (incl. nested bold+italic)
  - [text](url) links -> plain text with URL
  - fenced code blocks -> monospace paragraphs
  - markdown tables -> Word tables (Table Grid style)
  - bullet (- / *) and numbered (1.) lists
  - --- horizontal rules -> skipped (thin separator)
Sets core properties: title, author.
"""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_LINE_SPACING

INLINE_RE = re.compile(
    r'(\*\*.+?\*\*|\*[^*\n]+?\*|`[^`]+?`|\[[^\]]+?\]\([^)]+?\))'
)


def add_runs(paragraph, text):
    """Add runs to a paragraph, honoring **bold**, *italic*, `code`, [text](url)."""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            inner = part[2:-2]
            run = paragraph.add_run(inner)
            run.bold = True
            # nested *italic* inside bold: strip asterisks, apply italic to whole
            if inner.startswith('*') and inner.endswith('*') and len(inner) > 2:
                run.text = inner[1:-1]
                run.italic = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        elif part.startswith('['):
            m = re.match(r'\[([^\]]+?)\]\(([^)]+?)\)', part)
            if m:
                label, url = m.group(1), m.group(2)
                run = paragraph.add_run(label)
                run.bold = False
                paragraph.add_run(' (')
                urun = paragraph.add_run(url)
                urun.font.name = 'Courier New'
                urun.font.size = Pt(9)
                urun.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                paragraph.add_run(')')
            else:
                paragraph.add_run(part)
        else:
            paragraph.add_run(part)


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(8.5)


def flush_table(doc, rows):
    """rows: list of list-of-cell-strings; first row is header."""
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncol)
    table.style = 'Table Grid'
    table.autofit = True
    for i, row in enumerate(rows):
        for j in range(ncol):
            cell_text = row[j] if j < len(row) else ''
            cell = table.rows[i].cells[j]
            p = cell.paragraphs[0]
            add_runs(p, cell_text)
            for r in p.runs:
                r.font.size = Pt(9)
            if i == 0:
                for r in p.runs:
                    r.bold = True


def convert(md_path, docx_path):
    doc = Document()
    # base style
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.read().split('\n')

    i = 0
    in_code = False
    code_buf = []
    table_buf = []

    def flush_table_buf():
        nonlocal table_buf
        if table_buf:
            rows = []
            for t in table_buf:
                cells = [c.strip() for c in t.strip().strip('|').split('|')]
                # skip markdown separator rows like |---|---|
                if all(re.fullmatch(r':?-{2,}:?', c) for c in cells):
                    continue
                rows.append(cells)
            flush_table(doc, rows)
            table_buf = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith('```'):
            if in_code:
                add_code_block(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                flush_table_buf()
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # table line
        if line.strip().startswith('|') and line.strip().endswith('|'):
            table_buf.append(line)
            i += 1
            continue
        flush_table_buf()

        stripped = line.strip()

        # horizontal rule
        if re.fullmatch(r'-{3,}', stripped):
            i += 1
            continue

        # headings
        m = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            # strip inline markers in headings
            clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            clean = re.sub(r'\*([^*]+)\*', r'\1', clean)
            clean = clean.replace('`', '')
            doc.add_heading(clean, level=level)
            i += 1
            continue

        # bullet list
        m = re.match(r'^[-*]\s+(.*)$', stripped)
        if m:
            p = doc.add_paragraph(style='List Bullet')
            add_runs(p, m.group(1))
            i += 1
            continue

        # numbered list
        m = re.match(r'^(\d+)\.\s+(.*)$', stripped)
        if m:
            p = doc.add_paragraph(style='List Number')
            add_runs(p, m.group(2))
            i += 1
            continue

        # normal paragraph
        if stripped:
            p = doc.add_paragraph()
            add_runs(p, stripped)
        i += 1

    flush_table_buf()

    # core properties
    props = doc.core_properties
    props.title = ('Why We Build Our Agent Core in Rust '
                   '(and What It Would Take to Go All the Way)')
    props.author = 'agentgraph project'
    props.subject = 'Blog post'
    props.created = None if props.created is None else props.created

    doc.save(docx_path)


if __name__ == '__main__':
    src = sys.argv[1]
    dst = sys.argv[2]
    convert(src, dst)
    print(f'Converted {src} -> {dst}')
